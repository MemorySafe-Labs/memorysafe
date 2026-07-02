#!/usr/bin/env python3
"""Generate MemorySafe sales one-pager (honest v14.2 claims)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parent / "MemorySafe_Sales_OnePager.docx"
NAVY = RGBColor(30, 58, 95)
CYAN = RGBColor(0, 160, 180)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.85)
    s.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MemorySafe Labs")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Governed Replay for Rare-Class Continual Learning")
    r.font.size = Pt(12)
    r.font.color.rgb = CYAN

    doc.add_paragraph(
        "MemorySafe is a drop-in replay buffer policy that decides what to protect, replay, and evict "
        "under fixed memory — using MVI (forgetting risk) and ProtectScore, not replay frequency alone."
    )

    doc.add_heading("One-line pitch", level=2)
    doc.add_paragraph(
        "When your model learns continuously, rare but critical cases disappear from memory. "
        "MemorySafe governs a bounded buffer so they stay recoverable — without rewriting your architecture."
    )

    doc.add_heading("What you integrate (SKU)", level=2)
    doc.add_paragraph("MemorySafeBufferV14 — canonical product path (not the separate Brain governance demo).")
    doc.add_paragraph("~20 lines: insert batch → sample replay → train → update risk EMA.")

    doc.add_heading("Proof (reproducible, June 2026)", level=2)
    doc.add_paragraph("Protocol: v14.2-pneumonia-5task-sota · PneumoniaMNIST · 5 tasks · 20:1 imbalance · 500-sample buffer · 10 seeds")
    add_table(
        doc,
        ["Method", "Combined AUPRC", "vs MemorySafe (p)"],
        [
            ["Reservoir replay", "0.663 ± 0.066", "0.017"],
            ["Loss-priority (GSS)", "0.678 ± 0.062", "0.0015"],
            ["MemorySafe v14.2", "0.706 ± 0.051", "—"],
        ],
    )
    doc.add_paragraph("10/10 seed wins vs reservoir. Claim scope: this protocol — not universal CL SOTA (e.g. DER++ on CIFAR-100).")

    doc.add_heading("Safe claims vs retire", level=2)
    add_table(
        doc,
        ["Say this", "Do not say (without footnote)"],
        [
            ["SOTA on our PneumoniaMNIST 5-task harness (p < 0.05)", "Beats every CL method on every benchmark"],
            ["Combined AUPRC 0.706 ± 0.051 (10 seeds)", "0.941 AUPRC (legacy per-task metric / old protocol)"],
            ["Governed bounded replay under rare-class pressure", "Four-layer stack required for integration"],
            ["Drop-in buffer; model-agnostic", "Must replace your training stack"],
        ],
    )

    doc.add_heading("90-day pilot offer", level=2)
    for item in [
        "Your data stream + fixed buffer budget (500–5k samples)",
        "MemorySafe vs your baseline (reservoir or current replay)",
        "Deliverables: AUPRC, rare-class recall, buffer audit log, 1-page ROI",
        "Paid PoC recommended ($15–50k) — scopes seriousness",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Contact", level=2)
    p = doc.add_paragraph()
    p.add_run("Carla Centeno · Founder & CEO · MemorySafe Labs\n").bold = True
    p.add_run("carla@memorysafe.ca · memorysafe.ca · Montreal\n")
    p.add_run("NVIDIA Inception · Provisional patent (Apr 2026)")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run(f"Generated {date.today().isoformat()} · Desktop/memorysafe_v14/")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
